"""Tests for pairings_docx: rendering the final pairings as a Word doc."""

from __future__ import annotations

from pathlib import Path

import pytest

from pairings_docx import (
    _clean_court_label,
    _court_line,
    _format_doc_date,
    _ordinal,
    render_final_docx,
)


TEMPLATE = Path(__file__).parent / "tmp" / "Thursday Social Tennis.docx"


def test_ordinal_basic():
    assert _ordinal(1) == "st"
    assert _ordinal(2) == "nd"
    assert _ordinal(3) == "rd"
    assert _ordinal(4) == "th"
    assert _ordinal(21) == "st"
    assert _ordinal(22) == "nd"


def test_ordinal_teens_all_th():
    assert _ordinal(11) == "th"
    assert _ordinal(12) == "th"
    assert _ordinal(13) == "th"


def test_format_doc_date():
    assert _format_doc_date("2026-04-30") == "30th April"
    assert _format_doc_date("2026-05-01") == "1st May"
    assert _format_doc_date("2026-05-11") == "11th May"


def test_clean_court_label_courtreserve_style():
    assert _clean_court_label("Court #3 - Floodlit") == "Court 3"
    assert _clean_court_label("Court #11 - Floodlit:") == "Court 11"
    assert _clean_court_label("Court #7") == "Court 7"


def test_clean_court_label_plain_numeric():
    assert _clean_court_label("4") == "Court 4"
    assert _clean_court_label("10") == "Court 10"


def test_clean_court_label_passthrough():
    assert _clean_court_label("AY1") == "AY1"
    assert _clean_court_label("Outdoor") == "Outdoor"
    assert _clean_court_label("") == ""


def test_court_line_doubles():
    court = {
        "court_label": "4",
        "mode": "doubles",
        "players": ["Geoff Chapman", "Mei T", "Andy P", "Sarah F"],
        "pairs": [["Geoff Chapman", "Mei T"], ["Andy P", "Sarah F"]],
    }
    display = {"Geoff Chapman": "Geoff C", "Mei T": "Mei"}
    assert _court_line(court, display) == (
        "Court 4: Geoff C & Mei v Andy P & Sarah F"
    )


def test_court_line_courtreserve_label_gets_cleaned():
    court = {
        "court_label": "Court #3 - Floodlit",
        "mode": "doubles",
        "players": ["W", "J", "L", "A"],
        "pairs": [["W", "J"], ["L", "A"]],
    }
    assert _court_line(court, {}) == "Court 3: W & J v L & A"


def test_court_line_singles():
    court = {
        "court_label": "5", "mode": "singles",
        "players": ["A B", "C D"], "pairs": [["A B", "C D"]],
    }
    assert _court_line(court, {"A B": "A", "C D": "C"}) == "Court 5: A v C"


@pytest.mark.skipif(not TEMPLATE.exists(), reason="template not committed")
def test_render_final_docx_smoke(tmp_path):
    from docx import Document

    plan = {
        "date": "2026-05-14",
        "display_names": {
            "Geoff Chapman": "Geoff C", "Mei T": "Mei",
            "Andy P": "Andy", "Sarah F": "Sarah",
            "David K": "David", "Jack T": "Jack",
            "Louise Clark": "Louise",
        },
        "rotations": [
            {
                "rotation_num": 1, "start_time": "19:30", "end_time": "20:15",
                "sit_outs": ["Louise Clark"],
                "courts": [
                    {
                        "court_label": "4", "mode": "doubles",
                        "players": ["Geoff Chapman", "Mei T", "Andy P", "Sarah F"],
                        "pairs": [["Geoff Chapman", "Mei T"], ["Andy P", "Sarah F"]],
                    },
                    {
                        "court_label": "5", "mode": "singles",
                        "players": ["David K", "Jack T"],
                        "pairs": [["David K", "Jack T"]],
                    },
                ],
            },
            {
                "rotation_num": 2, "start_time": "20:15", "end_time": "20:55",
                "sit_outs": [],
                "courts": [
                    {
                        "court_label": "4", "mode": "doubles",
                        "players": ["Geoff Chapman", "Sarah F", "Mei T", "Jack T"],
                        "pairs": [["Geoff Chapman", "Sarah F"], ["Mei T", "Jack T"]],
                    },
                ],
            },
        ],
        "notes": "",
    }
    out = tmp_path / "out.docx"
    result = render_final_docx(plan, TEMPLATE, out)
    assert result == out
    assert out.exists()

    d = Document(str(out))
    texts = [p.text for p in d.paragraphs]

    # P0 (instructions) is the template's first paragraph — unchanged.
    assert "Thursday Social Tennis" in texts[0]
    assert "QR code" in texts[0]
    # P1 is the QR-code image (empty text but the image is still there).
    assert len(d.inline_shapes) == 1
    # P2 is the date heading we rewrote.
    assert texts[2] == "Pairings for 14th May"
    # The rest is rotation content.
    body = "\n".join(texts[3:])
    assert "Rotation 1 (19:30-20:15)" in body
    assert "Court 4: Geoff C & Mei v Andy & Sarah" in body
    assert "Court 5: David v Jack" in body
    assert "Rotation 2 (20:15-20:55)" in body
    assert "Sitting out (rotated fairly): Louise" in body

    # Spacing: every rotation-block paragraph should have 2pt after
    # (not Word's default 8pt) so the doc fits on one page.
    from docx.shared import Pt
    rotation_paras = [p for p in d.paragraphs if p.text.startswith("Court ")]
    assert rotation_paras, "no Court lines found"
    for p in rotation_paras:
        assert p.paragraph_format.space_after == Pt(2), (
            f"{p.text!r} should have 2pt-after, got "
            f"{p.paragraph_format.space_after}"
        )
