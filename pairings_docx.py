"""Render a finalised pairing plan as a Word document.

Used to produce the printable / shareable "Thursday Social Tennis"
poster that goes alongside the WhatsApp message. The template at
``tmp/Thursday Social Tennis.docx`` carries the instructions block
and the QR-code image; this module keeps those intact and rewrites
the date + rotation blocks from the plan dict.

Public API:
    render_final_docx(plan_dict, template_path, output_path) -> Path
"""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt


HEADING_PT = 16   # "Pairings for ..." line
ROTATION_PT = 12  # rotation header AND court lines
# Tight spacing so the whole doc fits on one A4 page. Word's default
# Normal style has 8pt-after, which pushes a 3-rotation evening onto a
# second page.
SPACE_AFTER_PT = 2          # between paragraphs in the rotation block
SPACE_BEFORE_HEADER_PT = 6  # extra room before each rotation header
SPACE_AFTER_DATE_PT = 8     # gap below the "Pairings for ..." heading


def _ordinal(n: int) -> str:
    if 11 <= n % 100 <= 13:
        return "th"
    return {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")


def _format_doc_date(iso_date: str) -> str:
    d = date.fromisoformat(iso_date)
    return f"{d.day}{_ordinal(d.day)} {d.strftime('%B')}"


def _clean_court_label(raw: str) -> str:
    """Reduce a CourtReserve-style label to a clean "Court N" form.

    CourtReserve gives long labels like ``"Court #3 - Floodlit"`` or
    ``"Court #11 - Floodlit:"``; the printed poster reads better as
    just ``"Court 3"``. Logic:

      * If the label contains ``#`` followed by digits, use those:
        ``"Court #3 - Floodlit"`` → ``"Court 3"``.
      * Else if the label IS a number (e.g. ``"7"``), prefix
        ``"Court "``: ``"Court 7"``.
      * Else (e.g. ``"AY1"``, ``"Outdoor"``) return as-is — anything
        unusual stays verbatim so the doc isn't misleading.
    """
    s = (raw or "").strip().rstrip(":").strip()
    if "#" in s:
        after = s.split("#", 1)[1]
        m = re.match(r"\s*(\d+)", after)
        if m:
            return f"Court {m.group(1)}"
    if s.isdigit():
        return f"Court {s}"
    return s


def _court_line(court: dict, display_names: dict[str, str]) -> str:
    label = _clean_court_label(court["court_label"])
    pairs = court["pairs"]

    def short(name: str) -> str:
        return display_names.get(name, name)

    if court["mode"] == "doubles":
        pa, pb = pairs[0], pairs[1]
        return (
            f"{label}: {short(pa[0])} & {short(pa[1])}"
            f" v {short(pb[0])} & {short(pb[1])}"
        )
    if court["mode"] == "singles":
        a, b = pairs[0]
        return f"{label}: {short(a)} v {short(b)}"
    return f"{label}: ?"


def _set_spacing(
    para,
    *,
    space_after_pt: float = SPACE_AFTER_PT,
    space_before_pt: float = 0,
) -> None:
    pf = para.paragraph_format
    pf.space_after = Pt(space_after_pt)
    pf.space_before = Pt(space_before_pt)


def _replace_paragraph_text(
    p, new_text: str, *, pt: int, bold: bool,
    space_after_pt: float = SPACE_AFTER_PT,
    space_before_pt: float = 0,
) -> None:
    """Wipe a paragraph's runs and replace with one styled run."""
    # Drop all existing runs.
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run(new_text)
    run.font.size = Pt(pt)
    if bold:
        run.bold = True
    _set_spacing(p, space_after_pt=space_after_pt, space_before_pt=space_before_pt)


def _delete_paragraph(p) -> None:
    p._element.getparent().remove(p._element)


def _add_para(
    doc, text: str, *, pt: int, bold: bool,
    space_after_pt: float = SPACE_AFTER_PT,
    space_before_pt: float = 0,
):
    para = doc.add_paragraph()
    if text:
        run = para.add_run(text)
        run.font.size = Pt(pt)
        if bold:
            run.bold = True
    _set_spacing(
        para,
        space_after_pt=space_after_pt,
        space_before_pt=space_before_pt,
    )
    return para


def render_final_docx(
    plan: dict,
    template_path: str | Path,
    output_path: str | Path,
) -> Path:
    """Render ``plan`` as a Word document at ``output_path``.

    The template's first two paragraphs (instructions + QR image) are
    kept verbatim. The "Pairings for ..." heading is updated to the
    plan's date; all subsequent paragraphs are replaced with the
    rotation blocks. Returns the output path.
    """
    template_path = Path(template_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(template_path))

    paragraphs = list(doc.paragraphs)
    if len(paragraphs) < 3:
        raise ValueError(
            f"template at {template_path} has only {len(paragraphs)} "
            "paragraphs — expected ≥ 3 (instructions, QR, heading)"
        )

    # P0 = instructions, P1 = QR image, P2 = date heading. Keep 0+1
    # untouched; rewrite 2's text; delete everything after.
    date_text = f"Pairings for {_format_doc_date(plan['date'])}"
    _replace_paragraph_text(
        paragraphs[2], date_text, pt=HEADING_PT, bold=True,
        space_after_pt=SPACE_AFTER_DATE_PT,
    )
    for p in paragraphs[3:]:
        _delete_paragraph(p)

    display_names: dict[str, str] = plan.get("display_names") or {}

    rotations = plan.get("rotations") or []
    for ri, rot in enumerate(rotations):
        start = rot.get("start_time", "")
        end = rot.get("end_time", "")
        header = f"Rotation {rot.get('rotation_num', ri + 1)} ({start}-{end})"
        _add_para(
            doc, header, pt=ROTATION_PT, bold=True,
            space_before_pt=(SPACE_BEFORE_HEADER_PT if ri > 0 else 0),
        )
        for court in rot.get("courts") or []:
            _add_para(
                doc, _court_line(court, display_names),
                pt=ROTATION_PT, bold=False,
            )

    # Sit-outs (collected across rotations; usually identical names
    # cycling, so just list the set).
    sit_outs: list[str] = []
    seen: set[str] = set()
    for rot in rotations:
        for s in rot.get("sit_outs") or []:
            if s not in seen:
                seen.add(s)
                sit_outs.append(s)
    if sit_outs:
        short = display_names
        _add_para(
            doc,
            "Sitting out (rotated fairly): "
            + ", ".join(short.get(n, n) for n in sit_outs),
            pt=ROTATION_PT, bold=False,
            space_before_pt=SPACE_BEFORE_HEADER_PT,
        )

    notes = (plan.get("notes") or "").strip()
    if notes:
        _add_para(
            doc, notes, pt=ROTATION_PT, bold=False,
            space_before_pt=SPACE_BEFORE_HEADER_PT,
        )

    doc.save(str(output_path))
    return output_path
